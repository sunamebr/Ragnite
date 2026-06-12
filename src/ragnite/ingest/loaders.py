"""File and directory loaders.

Plain text formats are handled with the standard library. PDF and DOCX are
supported via optional extras (``ragnite[pdf]`` / ``ragnite[docx]``).
"""

from __future__ import annotations

import json
import re
from html.parser import HTMLParser
from pathlib import Path

from ragnite.errors import IngestionError, MissingDependencyError
from ragnite.types import Document

TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".csv",
    ".tsv",
    ".log",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".go",
    ".rs",
    ".java",
    ".kt",
    ".c",
    ".h",
    ".cpp",
    ".hpp",
    ".cs",
    ".rb",
    ".php",
    ".swift",
    ".sql",
    ".sh",
    ".ps1",
    ".bat",
    ".r",
    ".scala",
    ".lua",
    ".xml",
}

IGNORED_DIRS = {
    ".git",
    ".hg",
    ".svn",
    "node_modules",
    "__pycache__",
    ".venv",
    "venv",
    "dist",
    "build",
    "vendor",
    ".ragnite",
    ".claude",
    ".idea",
    ".vscode",
    "target",
    ".pytest_cache",
    ".ruff_cache",
    ".mypy_cache",
}


class _HTMLTextExtractor(HTMLParser):
    _SKIP = {"script", "style", "noscript", "head"}

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self._SKIP:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP and self._skip_depth > 0:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0 and data.strip():
            self._parts.append(data.strip())

    def text(self) -> str:
        return "\n".join(self._parts)


def _read_text(path: Path) -> str:
    # utf-8-sig: transparently strips a BOM when present (no-op otherwise)
    return path.read_text(encoding="utf-8-sig", errors="replace")


def _load_html(path: Path) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(_read_text(path))
    return parser.text()


def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise MissingDependencyError("pypdf", "pdf") from exc
    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _load_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise MissingDependencyError("python-docx", "docx") from exc
    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs if p.text.strip())


def _load_json(path: Path) -> str:
    data = json.loads(_read_text(path))
    return json.dumps(data, ensure_ascii=False, indent=2)


def load_text(text: str, source: str = "inline", metadata: dict | None = None) -> Document:
    """Wrap a raw string as a Document."""
    return Document(text=text, source=source, metadata=metadata or {})


def _load_file(path: Path) -> Document | None:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _load_pdf(path)
    elif suffix == ".docx":
        text = _load_docx(path)
    elif suffix in {".html", ".htm"}:
        text = _load_html(path)
    elif suffix == ".json":
        text = _load_json(path)
    elif suffix == ".jsonl":
        text = _read_text(path)
    elif suffix in TEXT_SUFFIXES:
        text = _read_text(path)
    else:
        return None
    text = re.sub(r"\n{4,}", "\n\n\n", text).strip()
    if not text:
        return None
    return Document(
        text=text,
        source=str(path),
        metadata={"filename": path.name, "suffix": suffix},
    )


def load_path(path: str | Path, recursive: bool = True) -> list[Document]:
    """Load a file or every supported file under a directory."""
    root = Path(path)
    if not root.exists():
        raise IngestionError(f"path does not exist: {root}")
    if root.is_file():
        doc = _load_file(root)
        return [doc] if doc else []

    pattern = "**/*" if recursive else "*"
    docs: list[Document] = []
    for candidate in sorted(root.glob(pattern)):
        if not candidate.is_file():
            continue
        if any(part in IGNORED_DIRS for part in candidate.parts):
            continue
        doc = _load_file(candidate)
        if doc:
            docs.append(doc)
    return docs
